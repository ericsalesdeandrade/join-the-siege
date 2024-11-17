import os
import random

import pandas as pd
from docx import Document
from faker import Faker
from PIL import Image, ImageDraw, ImageFont

faker = Faker()

FONT_PATH = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_SIZE = 14

try:
    font = ImageFont.truetype(FONT_PATH, FONT_SIZE)
except OSError:
    print("Custom font not found. Using default font.")
    font = ImageFont.load_default()

TEMPLATE_SIZE = (800, 600)
BACKGROUND_COLOR = (255, 255, 255)
TEXT_COLOR = (0, 0, 0)


def generate_unique_filename(output_dir, base_name, extension):
    """
    Generates a unique filename by appending a random number.
    """
    random_number = random.randint(1000, 9999)
    return os.path.join(output_dir, f"{base_name}_{random_number}.{extension}")


def generate_random_invoice_items():
    """
    Generate a list of random invoice items.
    """
    items = []
    for _ in range(random.randint(3, 10)):
        description = random.choice(
            [
                "Laptop",
                "Monitor",
                "Keyboard",
                "Mouse",
                "Desk Chair",
                "External Hard Drive",
                "Headphones",
                "Webcam",
                "Printer",
            ]
        )
        quantity = random.randint(1, 5)
        unit_price = round(random.uniform(50, 500), 2)
        total_price = round(quantity * unit_price, 2)
        items.append(
            {
                "Description": description,
                "Quantity": quantity,
                "Unit Price": unit_price,
                "Total": total_price,
            }
        )
    return items


def calculate_invoice_totals(items):
    """
    Calculate subtotals, tax, and grand total.
    """
    subtotal = sum(item["Total"] for item in items)
    tax = round(subtotal * 0.1, 2)
    total = round(subtotal + tax, 2)
    return subtotal, tax, total


def generate_invoice_as_pdf(output_dir, invoice_data):
    """
    Generate invoices as PDFs.
    """
    os.makedirs(output_dir, exist_ok=True)
    image = Image.new("RGB", TEMPLATE_SIZE, color=BACKGROUND_COLOR)
    draw = ImageDraw.Draw(image)

    draw.text(
        (20, 20),
        f"Invoice Number: {invoice_data['Invoice Number']}",
        fill=TEXT_COLOR,
        font=font,
    )
    draw.text((20, 50), f"Date: {invoice_data['Date']}", fill=TEXT_COLOR, font=font)
    draw.text(
        (20, 80), f"Company: {invoice_data['Company']}", fill=TEXT_COLOR, font=font
    )
    draw.text(
        (20, 110), f"Customer: {invoice_data['Customer']}", fill=TEXT_COLOR, font=font
    )
    draw.text(
        (20, 140), f"Address: {invoice_data['Address']}", fill=TEXT_COLOR, font=font
    )

    y_offset = 180
    draw.text(
        (20, y_offset),
        "Description | Quantity | Unit Price | Total",
        fill=TEXT_COLOR,
        font=font,
    )
    y_offset += 20
    for item in invoice_data["Items"]:
        draw.text(
            (20, y_offset),
            f"{item['Description']} | {item['Quantity']} | ${item['Unit Price']:.2f} | ${item['Total']:.2f}",
            fill=TEXT_COLOR,
            font=font,
        )
        y_offset += 20

    draw.text(
        (20, y_offset + 20),
        f"Subtotal: ${invoice_data['Subtotal']:.2f}",
        fill=TEXT_COLOR,
        font=font,
    )
    draw.text(
        (20, y_offset + 40),
        f"Tax (10%): ${invoice_data['Tax']:.2f}",
        fill=TEXT_COLOR,
        font=font,
    )
    draw.text(
        (20, y_offset + 60),
        f"Total: ${invoice_data['Total']:.2f}",
        fill=TEXT_COLOR,
        font=font,
    )

    pdf_file_path = generate_unique_filename(output_dir, "invoice", "pdf")
    image.save(pdf_file_path, format="PDF")
    print(f"Saved: {pdf_file_path}")


def generate_invoice_as_docx(output_dir, invoice_data):
    """
    Generate invoices as DOCX files.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Invoice", level=1)

    doc.add_paragraph(f"Invoice Number: {invoice_data['Invoice Number']}")
    doc.add_paragraph(f"Date: {invoice_data['Date']}")
    doc.add_paragraph(f"Company: {invoice_data['Company']}")
    doc.add_paragraph(f"Customer: {invoice_data['Customer']}")
    doc.add_paragraph(f"Address: {invoice_data['Address']}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Quantity"
    hdr_cells[2].text = "Unit Price"
    hdr_cells[3].text = "Total"
    for item in invoice_data["Items"]:
        row_cells = table.add_row().cells
        row_cells[0].text = item["Description"]
        row_cells[1].text = str(item["Quantity"])
        row_cells[2].text = f"${item['Unit Price']:.2f}"
        row_cells[3].text = f"${item['Total']:.2f}"

    doc.add_paragraph(f"Subtotal: ${invoice_data['Subtotal']:.2f}")
    doc.add_paragraph(f"Tax (10%): ${invoice_data['Tax']:.2f}")
    doc.add_paragraph(f"Total: ${invoice_data['Total']:.2f}")

    docx_file_path = generate_unique_filename(output_dir, "invoice", "docx")
    doc.save(docx_file_path)
    print(f"Saved: {docx_file_path}")


def generate_invoice_as_xlsx(output_dir, invoice_data):
    """
    Generate invoices as XLSX files.
    """
    os.makedirs(output_dir, exist_ok=True)
    df = pd.DataFrame(invoice_data["Items"])
    xlsx_file_path = generate_unique_filename(output_dir, "invoice", "xlsx")
    df.to_excel(xlsx_file_path, index=False)
    print(f"Saved: {xlsx_file_path}")


def generate_synthetic_invoices(output_dir, format_type, count=10):
    """
    Generate synthetic invoices in the specified format.
    """
    for _ in range(count):
        items = generate_random_invoice_items()
        subtotal, tax, total = calculate_invoice_totals(items)
        invoice_data = {
            "Invoice Number": f"INV-{random.randint(1000, 9999)}",
            "Date": faker.date_between(start_date="-1y", end_date="today").strftime(
                "%m/%d/%Y"
            ),
            "Company": faker.company(),
            "Customer": faker.name(),
            "Address": faker.address().replace("\n", ", "),
            "Items": items,
            "Subtotal": subtotal,
            "Tax": tax,
            "Total": total,
        }

        if format_type == "pdf":
            generate_invoice_as_pdf(output_dir, invoice_data)
        elif format_type == "docx":
            generate_invoice_as_docx(output_dir, invoice_data)
        elif format_type == "xlsx":
            generate_invoice_as_xlsx(output_dir, invoice_data)


if __name__ == "__main__":
    output_dir = "./training_data/invoices/"
    generate_synthetic_invoices(output_dir, "pdf", count=900)
    generate_synthetic_invoices(output_dir, "docx", count=50)
    generate_synthetic_invoices(output_dir, "xlsx", count=50)
