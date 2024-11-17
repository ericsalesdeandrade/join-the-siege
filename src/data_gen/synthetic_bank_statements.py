import os
import random

import pandas as pd
from docx import Document
from faker import Faker
from PIL import Image, ImageDraw, ImageFont

faker = Faker()

FONT_PATH = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_SIZE = 14

try:
    font = ImageFont.truetype(FONT_PATH, FONT_SIZE)
except OSError:
    print("Custom font not found. Using default font.")
    font = ImageFont.load_default()

TEMPLATE_SIZE = (800, 600)
BACKGROUND_COLOR = (255, 255, 255)
TEXT_COLOR = (0, 0, 0)


def generate_unique_filename(output_dir, base_name, extension):
    """
    Generates a unique filename by appending a random number.
    """
    random_number = random.randint(1000, 9999)
    return os.path.join(output_dir, f"{base_name}_{random_number}.{extension}")


def generate_random_transactions():
    """
    Generate a list of random transactions based on provided examples.
    """
    transactions = []
    for _ in range(random.randint(10, 20)):
        date = faker.date_between(start_date="-1y", end_date="today").strftime(
            "%d/%m/%Y"
        )
        description = random.choice(
            [
                "Direct Deposit",
                "POS Purchase",
                "ATM Withdrawal",
                "Wire Transfer",
                "Loan Repayment",
                "ACH Payment",
                "Bank Fee",
                "Check Deposit",
                "Interest Credit",
            ]
        )
        debit = round(random.uniform(10, 1000), 2) if random.random() < 0.7 else ""
        credit = round(random.uniform(10, 1000), 2) if debit == "" else ""
        transactions.append(
            {
                "Date": date,
                "Description": description,
                "Debit ($)": debit,
                "Credit ($)": credit,
            }
        )
    return transactions


def generate_bank_statement_as_pdf(output_dir, transactions):
    """
    Generate bank statements as PDFs.
    """
    os.makedirs(output_dir, exist_ok=True)
    image = Image.new("RGB", TEMPLATE_SIZE, color=BACKGROUND_COLOR)
    draw = ImageDraw.Draw(image)

    draw.text(
        (20, 20), f"Bank Name: Bank of {faker.city()}", fill=TEXT_COLOR, font=font
    )
    draw.text((20, 50), f"Account Holder: {faker.name()}", fill=TEXT_COLOR, font=font)
    draw.text(
        (20, 80),
        f"Account Number: XXXX-XXXX-XXXX-{random.randint(1000, 9999)}",
        fill=TEXT_COLOR,
        font=font,
    )
    draw.text(
        (20, 110),
        f"Statement Period: {faker.date_this_month().strftime('%B %Y')}",
        fill=TEXT_COLOR,
        font=font,
    )
    draw.text(
        (20, 150),
        "Date | Description | Debit ($) | Credit ($)",
        fill=TEXT_COLOR,
        font=font,
    )

    y_offset = 180
    for transaction in transactions:
        transaction_line = f"{transaction['Date']} | {transaction['Description'][:20]} | {transaction['Debit ($)']} | {transaction['Credit ($)']}"
        draw.text((20, y_offset), transaction_line, fill=TEXT_COLOR, font=font)
        y_offset += 20

    pdf_file_path = generate_unique_filename(output_dir, "bank_statement", "pdf")
    image.save(pdf_file_path, format="PDF")
    print(f"Saved: {pdf_file_path}")


def generate_bank_statement_as_docx(output_dir, transactions):
    """
    Generate bank statements as DOCX files.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Bank Statement", level=1)

    doc.add_paragraph(f"Bank Name: Bank of {faker.city()}")
    doc.add_paragraph(f"Account Holder: {faker.name()}")
    doc.add_paragraph(f"Account Number: XXXX-XXXX-XXXX-{random.randint(1000, 9999)}")
    doc.add_paragraph(f"Statement Period: {faker.date_this_month().strftime('%B %Y')}")
    doc.add_paragraph("Date | Description | Debit ($) | Credit ($)")

    for transaction in transactions:
        transaction_line = f"{transaction['Date']} | {transaction['Description']} | Debit: {transaction['Debit ($)']} | Credit: {transaction['Credit ($)']}"
        doc.add_paragraph(transaction_line)

    docx_file_path = generate_unique_filename(output_dir, "bank_statement", "docx")
    doc.save(docx_file_path)
    print(f"Saved: {docx_file_path}")


def generate_bank_statement_as_xlsx(output_dir, transactions):
    """
    Generate bank statements as XLSX files.
    """
    os.makedirs(output_dir, exist_ok=True)
    df = pd.DataFrame(transactions)
    xlsx_file_path = generate_unique_filename(output_dir, "bank_statement", "xlsx")
    df.to_excel(xlsx_file_path, index=False)
    print(f"Saved: {xlsx_file_path}")


def generate_synthetic_bank_statements(output_dir, format_type, count=10):
    """
    Generate synthetic bank statements in the specified format.
    """
    for _ in range(count):
        transactions = generate_random_transactions()
        if format_type == "pdf":
            generate_bank_statement_as_pdf(output_dir, transactions)
        elif format_type == "docx":
            generate_bank_statement_as_docx(output_dir, transactions)
        elif format_type == "xlsx":
            generate_bank_statement_as_xlsx(output_dir, transactions)


if __name__ == "__main__":
    output_base_dir = "./training_data/bank_statements/"
    generate_synthetic_bank_statements(output_base_dir, "pdf", count=900)
    generate_synthetic_bank_statements(output_base_dir, "docx", count=50)
    generate_synthetic_bank_statements(output_base_dir, "xlsx", count=50)
